from collections import defaultdict
from docx import Document
import os

path = r"C:\Users\guangqiiang.lu\Pictures\tmp"
file_name = "节能报告基本情况.docx"
file_path = os.path.join(path, file_name)

word = Document(file_path)

tables = word.tables

def get_table_context(table):
    rows = table.rows

    key_names = [c.text for c in rows[0].cells]
    res = []
    
    for key in key_names:
        t = defaultdict()
        t[key] = []
        res.append(t)

    for i in range(1, len(rows)):
        for j in range(len(rows[i].cells)):
            res[j][key_names[j]].append(rows[i].cells[j].text)
    return res        


for table in word.tables:
    table_content = get_table_context(table)
    
    
res = [x.text for x in list(word.paragraphs) if x.text != ""]

print("*" * 20 )
print('\n'.join(res))


def get_row_dic(row):
    t = [x.strip() for x in row.split("：")]
    out = [m  for x in t for m in x.split(" ") if m!= ""]

    out_dic = {}
    for i in range(len(out)):
        if i % 2 == 0:
            out_dic[out[i]] = out[i+1]
        else:
            pass
        
    return out_dic
    

out_list = []

# this could be changed
sp_split = "："

for i in range(len(res)):
    if sp_split in res[i]:
        if res[i].count(sp_split)  == 1:
            sp = res[i].split(sp_split)
            
            out_list.append({sp[0]: sp[1]})
        elif res[i].count(sp_split) >1:
            out_list.append(get_row_dic(res[i]))
        else:
            pass
            
print("*"* 30)
print("Get final output : ")
print(out_list)


# This is for combine the structured information into a DataFrame
# make back into DF
import pandas as pd

def convert_table_to_df(table):
    r = [t.items() for t in table]

    df = pd.DataFrame()
    for i in range(len(r)):
        df = pd.concat([df, pd.DataFrame(dict(r[i]))], axis=1)
    
    return df
    
df = convert_table_to_df(table=table_content)

df.head()


import implicit

model = implicit.als.AlternatingLeastSquares()

import implicit.als


from sklearn.ensemble import GradientBoostingClassifier


from surprise import Dataset, evaluate

from surprise import NMF
